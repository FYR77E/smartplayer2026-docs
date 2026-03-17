#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from collections import OrderedDict
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml.ns import qn
from PIL import Image


ROOT_DIR = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT_DIR / "docs-source.docx"
DEFAULT_OUTPUT_DIR = ROOT_DIR / "docs" / "generated"
DEFAULT_ASSETS_DIR = ROOT_DIR / "static" / "img" / "manual"
WEBP_SOURCE_EXTENSIONS = {".png", ".jpg", ".jpeg"}

MAJOR_RE = re.compile(r"^(?P<number>\d+)\.\s+(?P<title>.+)$")
SUB_RE = re.compile(r"^(?P<number>\d+\.\d+)\.\s+(?P<title>.+)$")
FIGURE_RE = re.compile(r"^Рисунок\s+\d+\s*[—-].+$")
LEADING_NUMBER_RE = re.compile(r"^\d+\.\s*")
HEADING_BODY_STARTERS = (
    "Рисунок",
    "В",
    "После",
    "Перед",
    "Для",
    "При",
    "Если",
    "Данный",
    "Общий",
    "Симптом",
    "Решение",
    "Рекомендуется",
)


PREFACE_DOCS = [
    ("guide", "01-smartplayer-руководство.mdx", "SmartPlayer — Руководство", 1),
    ("user", "02-руководство-пользователя.mdx", "Руководство пользователя", 2),
    ("quickstart", "03-быстрый-старт.mdx", "Быстрый старт", 3),
    ("history", "04-история-изменений.mdx", "История изменений", 4),
    ("toc", "05-оглавление.mdx", "Оглавление", 5),
]

SECTION_DOCS = {
    "1": ("06-1-введение.mdx", "1. Введение", 6),
    "2": ("07-2-архитектура-и-развертывание.mdx", "2. Архитектура и развертывание", 7),
    "3": ("08-3-начало-работы.mdx", "3. Начало работы", 8),
    "4": ("09-4-разделы-мониторинга-и-устройств.mdx", "4. Разделы мониторинга и устройств", 9),
    "5": ("10-5-работа-с-контентом.mdx", "5. Работа с контентом", 10),
    "6": ("10-6-работа-с-трансляциями.mdx", "6. Работа с трансляциями", 11),
    "7": ("11-7-расписания.mdx", "7. Расписания", 12),
    "8": ("12-8-отчеты.mdx", "8. Отчеты", 13),
    "9": ("13-9-дополнительные-разделы.mdx", "9. Дополнительные разделы", 14),
    "10": ("14-10-администрирование.mdx", "10. Администрирование", 15),
    "11": ("15-11-практические-рекомендации.mdx", "11. Практические рекомендации", 16),
    "12": ("16-12-типовые-ошибки-и-решения.mdx", "12. Типовые ошибки и решения", 17),
    "13": ("17-13-чек-лист-запуска.mdx", "13. Чек-лист запуска", 18),
}


@dataclass
class ParaEntry:
    idx: int
    style: str
    text: str
    image_rel_ids: list[str]


@dataclass
class HeadingSplit:
    level: int
    heading: str
    remainder: str = ""


@dataclass
class ImageRef:
    rel_id: str
    source_name: str
    target_name: str
    size: int


@dataclass
class BlockImage:
    image_ref: ImageRef
    caption: str | None = None
    alt: str | None = None


@dataclass
class DocFile:
    filename: str
    title: str
    sidebar_position: int
    blocks: list[tuple[str, object]] = field(default_factory=list)
    _list_type: str | None = None
    _list_items: list[str] = field(default_factory=list)

    def flush_list(self) -> None:
        if self._list_type and self._list_items:
            self.blocks.append((self._list_type, list(self._list_items)))
        self._list_type = None
        self._list_items = []

    def add_heading(self, level: int, text: str) -> None:
        self.flush_list()
        self.blocks.append((f"h{level}", text))

    def add_paragraph(self, text: str) -> None:
        text = normalize_text(text)
        if not text:
            return
        self.flush_list()
        self.blocks.append(("p", text))

    def add_list_item(self, list_type: str, text: str) -> None:
        text = normalize_text(text)
        if not text:
            return
        if self._list_type != list_type:
            self.flush_list()
            self._list_type = list_type
        self._list_items.append(text)

    def add_image(self, image_ref: ImageRef, caption: str | None = None, alt: str | None = None) -> None:
        self.flush_list()
        self.blocks.append(("img", BlockImage(image_ref=image_ref, caption=caption, alt=alt)))

    def render(self, asset_prefix: str) -> str:
        lines = [
            "---",
            f'title: "{self.title}"',
            f"sidebar_position: {self.sidebar_position}",
            "---",
            "",
        ]

        for kind, value in self.blocks:
            if kind == "p":
                lines.extend([value, ""])
            elif kind == "ul":
                lines.extend([f"- {item}" for item in value])
                lines.append("")
            elif kind == "ol":
                lines.extend([f"1. {item}" for item in value])
                lines.append("")
            elif kind.startswith("h"):
                level = int(kind[1:])
                lines.extend([f'{"#" * level} {value}', ""])
            elif kind == "img":
                image: BlockImage = value
                alt = normalize_text(image.alt or image.caption or self.title)
                image_path = f"{asset_prefix}/{image.image_ref.target_name}"
                lines.extend([f"![{alt}]({image_path})", ""])
                if image.caption:
                    lines.extend([f"*{normalize_text(image.caption)}*", ""])

        while lines and not lines[-1]:
            lines.pop()
        return "\n".join(lines) + "\n"


def normalize_text(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\u200b", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def strip_number_prefix(text: str) -> str:
    return normalize_text(LEADING_NUMBER_RE.sub("", text, count=1))


def paragraph_images(document: Document, paragraph) -> list[ImageRef]:
    images: list[ImageRef] = []
    for run in paragraph.runs:
        for blip in run._element.findall(".//" + qn("a:blip")):
            rel_id = blip.get(qn("r:embed"))
            if not rel_id or rel_id not in document.part.rels:
                continue
            rel = document.part.rels[rel_id]
            source_name = Path(rel.target_ref).name
            source_path = Path(source_name)
            target_name = (
                f"{source_path.stem}.webp"
                if source_path.suffix.lower() in WEBP_SOURCE_EXTENSIONS
                else source_name
            )
            size = len(rel.target_part.blob)
            images.append(
                ImageRef(
                    rel_id=rel_id,
                    source_name=source_name,
                    target_name=target_name,
                    size=size,
                )
            )
    return images


def load_entries(source: Path) -> tuple[Document, list[ParaEntry]]:
    document = Document(str(source))
    entries: list[ParaEntry] = []
    for idx, paragraph in enumerate(document.paragraphs, start=1):
        text = normalize_text(paragraph.text)
        images = paragraph_images(document, paragraph)
        if not text and not images:
            continue
        entries.append(
            ParaEntry(
                idx=idx,
                style=paragraph.style.name,
                text=text,
                image_rel_ids=[image.rel_id for image in images],
            )
        )
    return document, entries


def find_toc_bounds(entries: list[ParaEntry]) -> tuple[int, int]:
    toc_start = next(i for i, entry in enumerate(entries) if entry.text == "Оглавление")
    body_start = next(
        i
        for i, entry in enumerate(entries[toc_start + 1 :], start=toc_start + 1)
        if entry.style == "Heading 1" and entry.text.startswith("1. ")
    )
    return toc_start, body_start


def toc_headings(entries: list[ParaEntry], toc_start: int, body_start: int) -> tuple[list[str], list[str], list[str]]:
    major: list[str] = []
    sub: list[str] = []
    toc_lines: list[str] = []
    for entry in entries[toc_start + 1 : body_start]:
        if not entry.text:
            continue
        toc_lines.append(entry.text)
        if MAJOR_RE.match(entry.text):
            major.append(entry.text)
        elif SUB_RE.match(entry.text):
            sub.append(entry.text)
    return toc_lines, major, sub


def longest_heading_prefix(text: str, headings: Iterable[str]) -> str | None:
    for heading in sorted(headings, key=len, reverse=True):
        if text == heading:
            return heading
        if text.startswith(heading + " "):
            return heading
    return None


def split_subheading_with_body(text: str) -> HeadingSplit | None:
    if not SUB_RE.match(text):
        return None

    starter_group = "|".join(re.escape(word) for word in HEADING_BODY_STARTERS)

    # Some DOCX paragraphs merge a short subsection title and the first sentence.
    quoted_match = re.match(
        rf"^(?P<heading>\d+\.\d+\.\s+.+?[»\"])(?=\s+(?:{starter_group})\b)(?P<rest>\s+.+)$",
        text,
    )
    if quoted_match:
        return HeadingSplit(
            level=2,
            heading=normalize_text(quoted_match.group("heading")),
            remainder=normalize_text(quoted_match.group("rest")),
        )

    generic_match = re.match(
        rf"^(?P<number>\d+\.\d+)\.\s+(?P<title>.+?)(?=\s+(?:{starter_group})\b)(?P<rest>\s+.+)$",
        text,
    )
    if not generic_match:
        return None

    title = normalize_text(generic_match.group("title"))
    if not 1 <= len(title.split()) <= 5:
        return None

    return HeadingSplit(
        level=2,
        heading=f"{generic_match.group('number')}. {title}",
        remainder=normalize_text(generic_match.group("rest")),
    )


def detect_heading_split(text: str, style: str, major_headings: list[str], sub_headings: list[str]) -> HeadingSplit | None:
    if not text:
        return None

    major_match = longest_heading_prefix(text, major_headings)
    if major_match:
        return HeadingSplit(level=1, heading=major_match, remainder=normalize_text(text[len(major_match) :]))

    sub_match = longest_heading_prefix(text, sub_headings)
    if sub_match:
        return HeadingSplit(level=2, heading=sub_match, remainder=normalize_text(text[len(sub_match) :]))

    if not style.startswith("List"):
        smart_sub_match = split_subheading_with_body(text)
        if smart_sub_match:
            return smart_sub_match

    if style == "Heading 1" and MAJOR_RE.match(text):
        return HeadingSplit(level=1, heading=text)
    if style == "Heading 2" and SUB_RE.match(text):
        return HeadingSplit(level=2, heading=text)

    if not style.startswith("List") and SUB_RE.match(text):
        return HeadingSplit(level=2, heading=text)

    if not style.startswith("List") and MAJOR_RE.match(text):
        return HeadingSplit(level=1, heading=text)

    return None


def build_preface_docs(entries: list[ParaEntry], toc_start: int, body_start: int, toc_lines: list[str]) -> dict[str, DocFile]:
    docs = {
        key: DocFile(filename=filename, title=title, sidebar_position=position)
        for key, filename, title, position in PREFACE_DOCS
    }

    for entry in entries[:toc_start]:
        if entry.text in {"SmartPlayer", "Руководство пользователя", "Быстрый старт", "История изменений"}:
            continue
        if 1 <= entry.idx <= 3:
            docs["guide"].add_paragraph(entry.text)
        if entry.idx == 3:
            docs["user"].add_paragraph(entry.text)
        if 5 <= entry.idx <= 11:
            if entry.style == "List Number":
                docs["quickstart"].add_list_item("ol", strip_number_prefix(entry.text))
            else:
                docs["quickstart"].add_paragraph(entry.text)
        if 13 <= entry.idx <= 17:
            docs["history"].add_list_item("ul", entry.text)

    current_indent = 0
    for line in toc_lines:
        if SUB_RE.match(line):
            docs["toc"].add_list_item("ul", f"  {line}")
        else:
            docs["toc"].add_list_item("ul", line)
    docs["guide"].add_paragraph("Документация SmartPlayer: пользовательское руководство, быстрый старт и структура разделов.")
    return docs


def extract_image_ref(document: Document, rel_id: str) -> ImageRef:
    rel = document.part.rels[rel_id]
    source_name = Path(rel.target_ref).name
    source_path = Path(source_name)
    target_name = (
        f"{source_path.stem}.webp"
        if source_path.suffix.lower() in WEBP_SOURCE_EXTENSIONS
        else source_name
    )
    return ImageRef(
        rel_id=rel_id,
        source_name=source_name,
        target_name=target_name,
        size=len(rel.target_part.blob),
    )


def is_decorative(image_ref: ImageRef, caption: str | None, text: str) -> bool:
    return image_ref.size < 3000 and caption is None and bool(text)


def next_caption(entries: list[ParaEntry], start: int) -> tuple[str | None, int]:
    if start >= len(entries):
        return None, start
    candidate = entries[start]
    if candidate.image_rel_ids:
        return None, start
    if FIGURE_RE.match(candidate.text):
        return candidate.text, start + 1
    return None, start


def build_section_docs(document: Document, entries: list[ParaEntry], body_start: int, major_headings: list[str], sub_headings: list[str]) -> tuple[dict[str, DocFile], set[str]]:
    docs = {
        key: DocFile(filename=filename, title=title, sidebar_position=position)
        for key, (filename, title, position) in SECTION_DOCS.items()
    }

    used_targets: set[str] = set()
    current_key: str | None = None
    i = body_start
    while i < len(entries):
        entry = entries[i]
        text = entry.text
        images = [extract_image_ref(document, rel_id) for rel_id in entry.image_rel_ids]
        heading = detect_heading_split(text, entry.style, major_headings, sub_headings)

        if heading and heading.level == 1:
            current_key = MAJOR_RE.match(heading.heading).group("number")
            if current_key not in docs:
                raise ValueError(f"Не настроен файл для раздела {heading.heading}")
            text = heading.remainder
            heading = None

        if current_key is None:
            i += 1
            continue

        target_doc = docs[current_key]

        if heading and heading.level == 2:
            target_doc.add_heading(2, heading.heading)
            text = heading.remainder

        if images:
            image_caption: str | None = None
            skip_to = i + 1

            if text and FIGURE_RE.match(text):
                image_caption = text
                text = ""
            elif text and FIGURE_RE.match(text.strip()):
                image_caption = text.strip()
                text = ""
            elif FIGURE_RE.match(text):
                image_caption = text
                text = ""

            if not text and not image_caption:
                image_caption, skip_to = next_caption(entries, i + 1)

            if text:
                if entry.style == "List Number":
                    target_doc.add_list_item("ol", strip_number_prefix(text))
                elif entry.style in {"List Bullet", "List Paragraph"}:
                    target_doc.add_list_item("ul", text)
                else:
                    target_doc.add_paragraph(text)

            for idx, image in enumerate(images):
                if is_decorative(image, image_caption, text):
                    continue
                used_targets.add(image.target_name)
                target_doc.add_image(
                    image,
                    caption=image_caption if idx == 0 else None,
                    alt=image_caption or target_doc.title,
                )
            i = skip_to
            continue

        if text:
            if entry.style == "List Number":
                target_doc.add_list_item("ol", strip_number_prefix(text))
            elif entry.style in {"List Bullet", "List Paragraph"}:
                target_doc.add_list_item("ul", text)
            else:
                target_doc.add_paragraph(text)

        i += 1

    return docs, used_targets


def copy_used_assets(document: Document, used_targets: set[str], assets_dir: Path) -> None:
    assets_dir.mkdir(parents=True, exist_ok=True)
    for existing in assets_dir.glob("*"):
        if existing.is_file():
            existing.unlink()

    for rel in document.part.rels.values():
        if not hasattr(rel, "target_ref"):
            continue
        source_name = Path(rel.target_ref).name
        source_path = Path(source_name)
        target_name = (
            f"{source_path.stem}.webp"
            if source_path.suffix.lower() in WEBP_SOURCE_EXTENSIONS
            else source_name
        )
        if target_name not in used_targets:
            continue
        output_path = assets_dir / target_name
        if source_path.suffix.lower() in WEBP_SOURCE_EXTENSIONS:
            image = Image.open(BytesIO(rel.target_part.blob))
            image.save(output_path, "WEBP", lossless=True, method=6)
        else:
            output_path.write_bytes(rel.target_part.blob)


def write_docs(docs: dict[str, DocFile], output_dir: Path, asset_prefix: str) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    for doc in docs.values():
        doc.flush_list()
        (output_dir / doc.filename).write_text(doc.render(asset_prefix), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Import SmartPlayer DOCX manual into Docusaurus docs.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE, help="Path to source .docx file")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR, help="Directory for generated docs")
    parser.add_argument("--assets-dir", type=Path, default=DEFAULT_ASSETS_DIR, help="Directory for extracted images")
    parser.add_argument(
        "--asset-prefix",
        default="/img/manual",
        help="Public path prefix for extracted assets inside generated docs",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    source = args.source.expanduser().resolve()
    if not source.exists():
        raise SystemExit(f"Source file not found: {source}")

    document, entries = load_entries(source)
    toc_start, body_start = find_toc_bounds(entries)
    toc_lines, major_headings, sub_headings = toc_headings(entries, toc_start, body_start)

    docs = build_preface_docs(entries, toc_start, body_start, toc_lines)
    section_docs, used_targets = build_section_docs(document, entries, body_start, major_headings, sub_headings)
    docs.update(section_docs)

    copy_used_assets(document, used_targets, args.assets_dir)
    write_docs(docs, args.output_dir, args.asset_prefix)

    generated = sorted(doc.filename for doc in docs.values())
    print(f"Imported {source.name}")
    print(f"Generated {len(generated)} docs into {args.output_dir}")
    print(f"Copied {len(used_targets)} images into {args.assets_dir}")


if __name__ == "__main__":
    main()
